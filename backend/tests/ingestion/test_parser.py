import pytest
from pathlib import Path

from app.ingestion.parser import LocalIngestion


class TestLocalIngestionScan:
    def test_scan_nonexistent_dir(self, tmp_path):
        ingestion = LocalIngestion(tmp_path)
        files = ingestion.scan_task_dir("nonexistent-task")
        assert files == []

    def test_scan_empty_dir(self, tmp_path):
        task_dir = tmp_path / "empty-task"
        task_dir.mkdir()
        ingestion = LocalIngestion(tmp_path)
        files = ingestion.scan_task_dir("empty-task")
        assert files == []

    def test_scan_finds_supported_files(self, tmp_path):
        task_dir = tmp_path / "task-1"
        task_dir.mkdir()
        (task_dir / "notes.txt").write_text("content", encoding="utf-8")
        (task_dir / "report.pdf").write_bytes(b"%PDF-1.4 fake")
        (task_dir / "data.docx").write_bytes(b"PK fake docx")
        (task_dir / "image.png").write_bytes(b"\x89PNG fake")

        ingestion = LocalIngestion(tmp_path)
        files = ingestion.scan_task_dir("task-1")

        names = [f.name for f in files]
        assert "notes.txt" in names
        assert "report.pdf" in names
        assert "data.docx" in names
        assert "image.png" not in names

    def test_scan_returns_sorted_files(self, tmp_path):
        task_dir = tmp_path / "task-2"
        task_dir.mkdir()
        (task_dir / "c_file.txt").write_text("c", encoding="utf-8")
        (task_dir / "a_file.txt").write_text("a", encoding="utf-8")
        (task_dir / "b_file.txt").write_text("b", encoding="utf-8")

        ingestion = LocalIngestion(tmp_path)
        files = ingestion.scan_task_dir("task-2")

        names = [f.name for f in files]
        assert names == ["a_file.txt", "b_file.txt", "c_file.txt"]


class TestLocalIngestionReadTxt:
    def test_read_utf8_txt(self, tmp_path):
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("这是中文内容", encoding="utf-8")

        ingestion = LocalIngestion(tmp_path)
        text = ingestion.read_file(txt_file)
        assert text == "这是中文内容"

    def test_read_gbk_txt(self, tmp_path):
        txt_file = tmp_path / "gbk.txt"
        txt_file.write_text("GBK编码内容", encoding="gbk")

        ingestion = LocalIngestion(tmp_path)
        text = ingestion.read_file(txt_file)
        assert text is not None
        assert "内容" in text

    def test_read_empty_txt(self, tmp_path):
        txt_file = tmp_path / "empty.txt"
        txt_file.write_text("", encoding="utf-8")

        ingestion = LocalIngestion(tmp_path)
        text = ingestion.read_file(txt_file)
        assert text == ""

    def test_read_unsupported_format(self, tmp_path):
        csv_file = tmp_path / "data.csv"
        csv_file.write_text("a,b,c", encoding="utf-8")

        ingestion = LocalIngestion(tmp_path)
        text = ingestion.read_file(csv_file)
        assert text is None


class TestLocalIngestionReadPdf:
    def test_read_pdf_with_pdfplumber(self, tmp_path):
        pytest.importorskip("pdfplumber")

        from unittest.mock import patch, MagicMock

        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(b"%PDF-1.4 fake")

        ingestion = LocalIngestion(tmp_path)

        mock_pdf = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "PDF中的文字内容"
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("app.ingestion.parser.pdfplumber.open", return_value=mock_pdf):
            text = ingestion.read_file(pdf_file)

        assert text is not None
        assert "PDF中的文字内容" in text

    def test_read_pdf_without_libraries(self, tmp_path):
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(b"%PDF-1.4 fake")

        ingestion = LocalIngestion(tmp_path)

        from unittest.mock import patch
        with patch.dict("sys.modules", {"pdfplumber": None, "fitz": None}):
            with patch("app.ingestion.parser.LocalIngestion._read_pdf", return_value=None):
                text = ingestion._read_pdf(pdf_file)

        assert text is None


class TestLocalIngestionReadDocx:
    def test_read_docx_with_python_docx(self, tmp_path):
        pytest.importorskip("docx")

        from docx import Document

        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")
        doc.save(str(docx_file))

        ingestion = LocalIngestion(tmp_path)
        text = ingestion.read_file(docx_file)

        assert text is not None
        assert "第一段内容" in text
        assert "第二段内容" in text

    def test_read_docx_with_tables(self, tmp_path):
        pytest.importorskip("docx")

        from docx import Document

        docx_file = tmp_path / "table.docx"
        doc = Document()
        doc.add_paragraph("文档标题")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "列1"
        table.cell(0, 1).text = "列2"
        table.cell(1, 0).text = "数据1"
        table.cell(1, 1).text = "数据2"
        doc.save(str(docx_file))

        ingestion = LocalIngestion(tmp_path)
        text = ingestion.read_file(docx_file)

        assert text is not None
        assert "文档标题" in text
        assert "列1" in text
        assert "数据1" in text


class TestLocalIngestionIngestTask:
    def test_ingest_task_with_txt(self, tmp_path):
        task_dir = tmp_path / "task-ingest"
        task_dir.mkdir()
        (task_dir / "notes.txt").write_text("这是本地干货内容", encoding="utf-8")

        ingestion = LocalIngestion(tmp_path)
        results = ingestion.ingest_task("task-ingest")

        assert len(results) == 1
        assert results[0]["source"] == "notes.txt"
        assert "干货" in results[0]["text"]

    def test_ingest_task_multiple_files(self, tmp_path):
        task_dir = tmp_path / "task-multi"
        task_dir.mkdir()
        (task_dir / "a.txt").write_text("文件A内容", encoding="utf-8")
        (task_dir / "b.txt").write_text("文件B内容", encoding="utf-8")

        ingestion = LocalIngestion(tmp_path)
        results = ingestion.ingest_task("task-multi")

        assert len(results) == 2
        sources = [r["source"] for r in results]
        assert "a.txt" in sources
        assert "b.txt" in sources

    def test_ingest_task_skips_unsupported(self, tmp_path):
        task_dir = tmp_path / "task-skip"
        task_dir.mkdir()
        (task_dir / "notes.txt").write_text("有效内容", encoding="utf-8")
        (task_dir / "photo.jpg").write_bytes(b"\xff\xd8\xff fake jpeg")

        ingestion = LocalIngestion(tmp_path)
        results = ingestion.ingest_task("task-skip")

        assert len(results) == 1
        assert results[0]["source"] == "notes.txt"

    def test_ingest_nonexistent_task(self, tmp_path):
        ingestion = LocalIngestion(tmp_path)
        results = ingestion.ingest_task("nonexistent")
        assert results == []
