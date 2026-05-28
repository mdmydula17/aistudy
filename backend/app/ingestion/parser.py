import logging
from pathlib import Path
from typing import Optional, List

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class LocalIngestion:
    def __init__(self, base_dir: Path):
        self._base_dir = base_dir

    def scan_task_dir(self, task_id: str) -> List[Path]:
        task_dir = self._base_dir / task_id
        if not task_dir.exists():
            logger.info(f"[scan_task_dir] Directory not found: {task_dir}")
            return []

        files = sorted(
            f for f in task_dir.iterdir()
            if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
        )
        logger.info(f"[scan_task_dir] Found {len(files)} files in {task_dir}")
        return files

    def read_file(self, filepath: Path) -> Optional[str]:
        suffix = filepath.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            logger.warning(f"[read_file] Unsupported format: {suffix}")
            return None

        if suffix == ".txt":
            return self._read_txt(filepath)
        elif suffix == ".pdf":
            return self._read_pdf(filepath)
        elif suffix == ".docx":
            return self._read_docx(filepath)
        return None

    def ingest_task(self, task_id: str) -> List[dict]:
        files = self.scan_task_dir(task_id)
        results = []
        for f in files:
            try:
                text = self.read_file(f)
                if text:
                    results.append({
                        "source": f.name,
                        "text": text,
                    })
                else:
                    logger.warning(f"[ingest_task] No text extracted from {f.name}")
            except Exception as e:
                logger.warning(f"[ingest_task] Failed to process {f.name}: {e}")
        return results

    def _read_txt(self, filepath: Path) -> Optional[str]:
        try:
            return filepath.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return filepath.read_text(encoding="gbk")
        except Exception as e:
            logger.error(f"[_read_txt] Failed: {e}")
            return None

    def _read_pdf(self, filepath: Path) -> Optional[str]:
        try:
            import pdfplumber

            texts = []
            with pdfplumber.open(str(filepath)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        texts.append(page_text)
            return "\n".join(texts) if texts else None
        except ImportError:
            logger.warning("[_read_pdf] pdfplumber not installed, trying PyMuPDF")
        except Exception as e:
            logger.error(f"[_read_pdf] pdfplumber failed: {e}")

        try:
            import fitz

            doc = fitz.open(str(filepath))
            texts = [page.get_text() for page in doc]
            doc.close()
            return "\n".join(texts) if texts else None
        except ImportError:
            logger.error("[_read_pdf] Neither pdfplumber nor PyMuPDF available")
        except Exception as e:
            logger.error(f"[_read_pdf] PyMuPDF failed: {e}")

        return None

    def _read_docx(self, filepath: Path) -> Optional[str]:
        try:
            from docx import Document

            doc = Document(str(filepath))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" |"):
                        tables_text.append(row_text)

            all_text = paragraphs
            if tables_text:
                all_text.append("--- 表格内容 ---")
                all_text.extend(tables_text)

            return "\n".join(all_text) if all_text else None
        except ImportError:
            logger.error("[_read_docx] python-docx not installed")
        except Exception as e:
            logger.error(f"[_read_docx] Failed: {e}")

        return None
